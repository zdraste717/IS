from django.shortcuts import render, redirect
from datetime import datetime
import re
from django.views.decorators.http import require_POST
from django.http import HttpResponse
import json
from docx import Document
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from django.views.decorators.csrf import csrf_exempt   
from io import BytesIO
from django.http import JsonResponse
from django.db.models import Q
from .models import (
    Student, Scienes, Sport, Creation, VariousLevel, Publication,
    StudentGovernment, OtherAchiev, AddProgramm, Experience
)

def achiev_view(request):
    fullname = request.session.get('student_fullname')
    if not fullname:
        return redirect('login')
    
    student = Student.objects.filter(fullname=fullname).first()

    data = {
        'scienes': Scienes.objects.filter(fullname=fullname),
        'sport': Sport.objects.filter(fullname=fullname),
        'creation': Creation.objects.filter(fullname=fullname),
        'various_level': VariousLevel.objects.filter(fullname=fullname),
        'publication': Publication.objects.filter(fullname=fullname),
        'student_government': StudentGovernment.objects.filter(fullname=fullname),
        'other_achiev': OtherAchiev.objects.filter(fullname=fullname),
        'add_programm': AddProgramm.objects.filter(fullname=fullname),
        'experience': Experience.objects.filter(fullname=fullname)
    }

    # Проверяем — есть ли хоть одно достижение
    has_achievements = any(qs.exists() for qs in data.values())

    return render(request, 'achiev.html', {
        **data,
        'fullname': fullname,
        'student': student,
        'has_achievements': has_achievements
    })

def main_view(request):
    return render(request, 'main.html')

def login_view(request):
    if request.method == 'POST':
        last_name = request.POST.get('last_name', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        middle_name = request.POST.get('middle_name', '').strip()

        full_name = f"{last_name} {first_name} {middle_name}"
        request.session['student_fullname'] = full_name

        try:
            student = Student.objects.get(fullname=full_name)
            request.session['group_s'] = student.group_s
            request.session['faculty'] = student.faculty
        except Student.DoesNotExist:
            request.session['group_s'] = None
            request.session['faculty'] = None

        # 👉 Определяем, является ли пользователь администратором
        if full_name.lower() == "махалкина татьяна олеговна":
            request.session['is_admin'] = True
        else:
            request.session['is_admin'] = False

        return redirect('main')

    return render(request, 'index.html')


def normalize_date(date_str):
    
    #Преобразует строку даты в формат YYYYMMDD, игнорируя разделители
    #Поддерживаются форматы: ДД.ММ.ГГГГ, ДД-ММ-ГГГГ, ГГГГ-ММ-ДД, и др.
   
    if not date_str:
        return None

    date_str = date_str.strip()

    # Пробуем определить формат автоматически
    known_formats = [
        '%d.%m.%Y',
        '%d-%m-%Y',
        '%Y-%m-%d',
        '%Y.%m.%d',
        '%d/%m/%Y',
        '%Y/%m/%d',
    ]

    for fmt in known_formats:
        try:
            parsed = datetime.strptime(date_str[:10], fmt)
            return parsed.strftime('%Y%m%d')
        except ValueError:
            continue

    # Если не удалось — убираем всё, кроме цифр (например, 12022022)
    digits = re.sub(r'\D', '', date_str)
    if len(digits) == 8:
        if int(digits[:4]) > 1900:
            return digits  # ГГГГММДД
        else:
            return digits[4:] + digits[2:4] + digits[0:2]  # ДДММГГГГ → ГГГГММДД
    return None



def admin_panel_view(request):
    if not request.session.get('is_admin'):
        return redirect('main')

    type_filter = request.GET.get('type')
    search_query = request.GET.get('search', '').strip()
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')

    # Мапа моделей с полями для поиска и датой
    model_map = {
        'scienes': (Scienes, ['fullname', 'srw_name', 'customer', 'l_programm', 'p_name', 'f_date', 'result'], 's_date'),
        'sport': (Sport, ['fullname', 'sport_t', 'compet', 'note'], 'date_c'),
        'creation': (Creation, ['fullname', 'activity_t', 'part_fest', 'note'], 'date_f'),
        'various_level': (VariousLevel, ['fullname', 'event_t', 'level_e', 'name_e', 'venue', 'result'], 'date_e'),
        'publication': (Publication, ['fullname', 'coauthors', 'fname_work', 'form_w', 'public_t', 'article_inc'], 'output_d'),
        'student_government': (StudentGovernment, ['fullname', 'body_g', 'activity_t', 'note'], 'activity_per'),
        'other_achiev': (OtherAchiev, ['fullname', 'activity_t', 'achiev', 'note'], 'date_o'),
        'add_programm': (AddProgramm, ['fullname', 'name_p', 'education_t', 'hours', 'education_p', 'name_d'], 'terms'),
        'experience': (Experience, ['fullname', 'place_w', 'work_t', 'title_j', 'respons'], 'deadlines'),
    }

def admin_panel_view(request):
    if not request.session.get('is_admin'):
        return redirect('main')

    type_filter = request.GET.get('type')
    search_query = request.GET.get('search', '').strip()
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')

    model_map = {
        'scienes': (Scienes, ['fullname', 'srw_name', 'customer', 'l_programm', 'p_name', 'f_date', 'result'], 's_date'),
        'sport': (Sport, ['fullname', 'sport_t', 'compet', 'note'], 'date_c'),
        'creation': (Creation, ['fullname', 'activity_t', 'part_fest', 'note'], 'date_f'),
        'various_level': (VariousLevel, ['fullname', 'event_t', 'level_e', 'name_e', 'venue', 'result'], 'date_e'),
        'publication': (Publication, ['fullname', 'coauthors', 'fname_work', 'form_w', 'public_t', 'article_inc'], 'output_d'),
        'student_government': (StudentGovernment, ['fullname', 'body_g', 'activity_t', 'note'], 'activity_per'),
        'other_achiev': (OtherAchiev, ['fullname', 'activity_t', 'achiev', 'note'], 'date_o'),
        'add_programm': (AddProgramm, ['fullname', 'name_p', 'education_t', 'hours', 'education_p', 'name_d'], 'terms'),
        'experience': (Experience, ['fullname', 'place_w', 'work_t', 'title_j', 'respons'], 'deadlines'),
    }

    def apply_filters(queryset, fields, date_field):
        if search_query:
            q = Q()
            for f in fields:
                q |= Q(**{f"{f}__icontains": search_query})
            queryset = queryset.filter(q)

        if date_from and date_to:
            norm_from = normalize_date(date_from)
            norm_to = normalize_date(date_to)

            filtered = []
            for obj in queryset:
                raw_date = str(getattr(obj, date_field))[:10]
                norm_obj_date = normalize_date(raw_date)
                if norm_obj_date and norm_from <= norm_obj_date <= norm_to:
                    filtered.append(obj)
            return filtered

        return queryset

    context = {}

    if not type_filter or type_filter == 'all':
        for key, (model, fields, date_field) in model_map.items():
            queryset = model.objects.exclude(status="В обработке")
            context[key] = apply_filters(queryset, fields, date_field)
    else:
        model_info = model_map.get(type_filter)
        if model_info:
            model, fields, date_field = model_info
            queryset = model.objects.exclude(status="В обработке")
            context[type_filter] = apply_filters(queryset, fields, date_field)

    context.update({
        'type_filter': type_filter or 'all',
        'search_query': search_query,
        'date_from': date_from,
        'date_to': date_to,
        'is_admin': request.session.get('is_admin', False)
    })

    return render(request, 'admin_panel.html', context)



from datetime import datetime

def parse_date(date_str):
    try:
        return datetime.strptime(date_str.strip(), "%d.%m.%Y").date()
    except (ValueError, AttributeError):
        return None
    
@require_POST
def submit_application(request):
    try:
        is_admin = request.session.get('is_admin', False)

        if is_admin:
            lastname = request.POST.get('lastname', '').strip()
            firstname = request.POST.get('firstname', '').strip()
            middlename = request.POST.get('middlename', '').strip()
            fullname = f"{lastname} {firstname} {middlename}".strip()
            group_s = request.POST.get('group_s', '').strip()
            faculty = request.POST.get('faculty', '').strip()
        else:
            fullname = request.session.get('student_fullname')
            group_s = request.session.get('group_s')
            faculty = request.session.get('faculty')

        force_status = "Принято" if request.POST.get('force_accepted') == '1' else "В обработке"
        achievement_type = request.POST.get('achievementType')

        if not (fullname and group_s and faculty):
            return JsonResponse({
                'success': False,
                'error': 'Не удалось получить данные из сессии. Пожалуйста, войдите в систему повторно.'
            })

        if achievement_type == 'science':
            Scienes.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                srw_name=request.POST.get('srw_name', '').strip(),
                customer=request.POST.get('customer', '').strip(),
                l_programm=request.POST.get('programm_c', '').strip(),
                p_name=request.POST.get('name_proj', '').strip(),
                s_date=parse_date(request.POST.get('dateS_event')),
                f_date=parse_date(request.POST.get('dateF_event')),
                result=request.POST.get('result_science', '').strip(),
                status=force_status
            )

        elif achievement_type == 'sport':
            Sport.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                sport_t=request.POST.get('sport_t', '').strip(),
                compet=request.POST.get('compet', '').strip(),
                date_c=parse_date(request.POST.get('date_compet')),
                note=request.POST.get('note_compet', '').strip(),
                status=force_status
            )

        elif achievement_type == 'creative':
            Creation.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                activity_t=request.POST.get('activity_t_creative', '').strip(),
                part_fest=request.POST.get('part_fest', '').strip(),
                date_f=parse_date(request.POST.get('date_fest')),
                note=request.POST.get('note_fest', '').strip(),
                status=force_status
            )

        elif achievement_type == 'various':
            VariousLevel.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                event_t=request.POST.get('various_t', '').strip(),
                level_e=request.POST.get('level_v', '').strip(),
                name_e=request.POST.get('name_v', '').strip(),
                venue=request.POST.get('place_v', '').strip(),
                date_e=parse_date(request.POST.get('date_various')),
                result=request.POST.get('result_v', '').strip(),
                status=force_status
            )

        elif achievement_type == 'public':
            Publication.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                coauthors=request.POST.get('coauthors', '').strip(),
                fname_work=request.POST.get('fullname_w', '').strip(),
                output_d=request.POST.get('output', '').strip(),  # Строка, не дата
                form_w=request.POST.get('form_w', '').strip(),
                public_t=request.POST.get('public_t', '').strip(),
                article_inc=request.POST.get('include', '').strip(),
                status=force_status
            )

        elif achievement_type == 'government':
            StudentGovernment.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                body_g=request.POST.get('government_t', '').strip(),
                activity_t=request.POST.get('activity_t', '').strip(),
                activity_per=request.POST.get('per_g', '').strip(),  # Строка, не дата
                note=request.POST.get('note_g', '').strip(),
                status=force_status
            )

        elif achievement_type == 'other':
            OtherAchiev.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                activity_t=request.POST.get('activity_o', '').strip(),
                achiev=request.POST.get('other_a', '').strip(),
                date_o=parse_date(request.POST.get('date_other')),
                note=request.POST.get('note_other', '').strip(),
                status=force_status
            )

        elif achievement_type == 'addprogramm':
            AddProgramm.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                name_p=request.POST.get('name_prog', '').strip(),
                education_t=request.POST.get('programm_t', '').strip(),
                hours=request.POST.get('hours_p', '').strip(),
                terms=request.POST.get('date_compet', '').strip(),  # Строка
                education_p=request.POST.get('place_t', '').strip(),
                name_d=request.POST.get('document', '').strip(),
                status=force_status
            )

        elif achievement_type == 'experience':
            Experience.objects.create(
                fullname=fullname,
                group_s=group_s,
                faculty=faculty,
                place_w=request.POST.get('place_e', '').strip(),
                work_t=request.POST.get('exp_t', '').strip(),
                title_j=request.POST.get('position', '').strip(),
                deadlines=request.POST.get('per_e', '').strip(),  # Строка
                respons=request.POST.get('duty', '').strip(),
                status=force_status
            )

        else:
            return JsonResponse({'success': False, 'error': 'Неверный тип достижения'})

        return JsonResponse({'success': True})

    except Exception as e:
        import traceback
        traceback.print_exc()  # Покажет полную трассировку в консоли
        return JsonResponse({'success': False, 'error': f'Ошибка на сервере: {str(e)}'})


def admin_pending_view(request):
    if not request.session.get('is_admin'):
        return redirect('main')

    model_map = {
    'scienes': Scienes,
    'sport': Sport,
    'creation': Creation,
    'various_level': VariousLevel,
    'publication': Publication,
    'student_government': StudentGovernment,
    'other_achiev': OtherAchiev,
    'add_programm': AddProgramm,
    'experience': Experience,
}

    context = {}

    for key, model in model_map.items():
        context[key] = model.objects.filter(status="В обработке")

    return render(request, 'admin_pending.html', context)


from django.views.decorators.http import require_POST
from django.http import JsonResponse

@require_POST
def update_achievement_status(request):
    achiev_type = request.POST.get('type')
    achiev_id = request.POST.get('id')
    action = request.POST.get('action')  # 'accept', 'reject' или 'edit'

    model_map = {
        'scienes': Scienes,
        'sport': Sport,
        'creation': Creation,
        'various_level': VariousLevel,
        'publication': Publication,
        'student_government': StudentGovernment,
        'other_achiev': OtherAchiev,
        'add_programm': AddProgramm,
        'experience': Experience,
    }

    model = model_map.get(achiev_type)
    if not model:
        return JsonResponse({'success': False, 'error': 'Неверный тип'})

    try:
        instance = model.objects.get(id=achiev_id)

        if action == 'accept':
            instance.status = 'Принято'
            instance.save()
        elif action == 'reject':
            instance.delete()
        elif action == 'edit':
            for key in request.POST:
                if key not in ['id', 'type', 'action', 'csrfmiddlewaretoken']:
                    setattr(instance, key, request.POST.get(key))
            instance.save()
        else:
            return JsonResponse({'success': False, 'error': 'Неверное действие'})

        return JsonResponse({'success': True})

    except model.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Объект не найден'})


@csrf_exempt
def generate_report_docx(request):
    if request.method != 'POST':
        return HttpResponse("Метод не разрешён", status=405)

    try:
        data = json.loads(request.POST.get('report_data', '[]'))
    except json.JSONDecodeError:
        return HttpResponse("Ошибка при обработке данных", status=400)

    if not data:
        return HttpResponse("Нет данных для отчёта", status=400)

    doc = Document()
    doc.add_heading('Отчёт о достижениях студентов', 0)

    # Группируем по типу
    grouped = {}
    for item in data:
        type_ = item.get('type', 'Без категории')
        grouped.setdefault(type_, []).append(item)

    for type_, achievements in grouped.items():
        valid_achievements = []
        for ach in achievements:
            fullname = ach.get('fullname', '').strip()
            content_lines = [
                line.strip() for line in ach.get('content', [])
                if line.strip() and not line.startswith('ФИО:')
            ]
            if fullname or content_lines:
                valid_achievements.append((fullname, content_lines))

        if not valid_achievements:
            continue  # Ничего не выводим, если все достижения пустые

        doc.add_heading(type_, level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        hdr = table.rows[0].cells
        hdr[0].text = 'ФИО' 
        hdr[1].text = 'Информация'

        for fullname, content_lines in valid_achievements:
            row = table.add_row().cells
            row[0].text = fullname
            row[1].text = '\n'.join(content_lines)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=report.docx'
    return response

@csrf_exempt
def generate_report_excel(request):
    if request.method != 'POST':
        return HttpResponse("Метод не разрешён", status=405)

    try:
        data = json.loads(request.POST.get('report_data', '[]'))
    except json.JSONDecodeError:
        return HttpResponse("Ошибка при обработке данных", status=400)

    if not data:
        return HttpResponse("Нет данных для отчёта", status=400)

    wb = Workbook()
    wb.remove(wb.active)  # Удаляем дефолтный пустой лист

    grouped = {}
    for item in data:
        type_ = item.get('type', 'Без категории')
        grouped.setdefault(type_, []).append(item)

    for type_name, achievements in grouped.items():
        ws = wb.create_sheet(title=type_name[:31])  # max 31 символ

        # Заголовки
        all_lines = [a['content'] for a in achievements if a['content']]
        flat_fields = set()

        for lines in all_lines:
            for line in lines:
                if ':' in line:
                    key = line.split(':')[0].strip()
                    flat_fields.add(key)

        flat_fields.discard("ФИО")
        headers = ["ФИО"] + sorted(flat_fields)
        ws.append(headers)

        for ach in achievements:
            row = []
            content_map = {}
            for line in ach.get('content', []):
                if ':' in line:
                    key, val = line.split(':', 1)
                    content_map[key.strip()] = val.strip()
            row = [content_map.get(header, "") for header in headers]
            ws.append(row)

        # Автоширина
        for col_idx, _ in enumerate(headers, start=1):
            col_letter = get_column_letter(col_idx)
            max_len = max((len(str(ws.cell(row=i, column=col_idx).value)) for i in range(1, ws.max_row + 1)), default=10)
            ws.column_dimensions[col_letter].width = max_len + 2

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=report.xlsx'
    return response

def logout_view(request):
    request.session.flush()  
    return redirect('login')  
